import io
import xml.etree.ElementTree as ET
from typing import List, TYPE_CHECKING
from datetime import datetime
from docx import Document
from docx.shared import Pt


if TYPE_CHECKING:
    from reportlab.platypus import Table

from ..models import Task


class ExportService:
    """Service for exporting tasks in various formats."""

    @staticmethod
    def _get_language_tag(language: str) -> str:
        """Get language tag for code blocks."""
        return language.lower() if language else "text"

    @staticmethod
    def _get_validation_info(task: Task) -> tuple:
        """Get validation status icon and text for a task."""
        validation = task.validation_result or {}
        status = validation.get('status', 'unknown')
        if status == 'passed':
            return "✅", "PASSED"
        return "❌", "FAILED"

    @staticmethod
    def _format_task_metadata(task: Task) -> List[str]:
        """Format task metadata as markdown lines."""
        lines = [
            f"**Language:** {task.language}",
            f"**Concept:** {task.concept}",
            f"**Difficulty:** {task.difficulty}",
        ]
        if hasattr(task, 'template_name') and task.template_name:
            lines.append(f"**Template:** {task.template_name}")
        return lines

    @staticmethod
    def _format_task_sections(task: Task) -> List[str]:
        """Format optional task sections (examples, solution, tests, validation)."""
        sections = []

        if task.examples and task.examples.strip():
            sections.append("### Examples\n")
            sections.append(f"{task.examples}\n")

        if task.solution:
            lang = ExportService._get_language_tag(task.language)
            sections.append("### Reference Solution\n")
            sections.append(f"```{lang}\n{task.solution}\n```\n")

        if task.tests and task.tests.strip():
            lang = ExportService._get_language_tag(task.language)
            sections.append("### Tests\n")
            sections.append(f"```{lang}\n{task.tests}\n```\n")

        if hasattr(task, 'is_validated') and task.is_validated:
            status_icon, status_text = ExportService._get_validation_info(task)
            sections.append("### Validation\n")
            sections.append(f"{status_icon} **{status_text}** ({(task.validation_result or {}).get('status', 'unknown').upper()})\n")

        return sections

    @staticmethod
    def _format_task(task: Task, index: int) -> List[str]:
        """Format a single task as markdown."""
        lines = []
        lines.append(f"\n## Task {index}: {task.title}\n")
        lines.extend(ExportService._format_task_metadata(task))
        lines.append("\n---\n")
        lines.append("### Description\n")
        lines.append(f"{task.description}\n")
        lines.extend(ExportService._format_task_sections(task))
        lines.append("---\n")
        return lines

    @staticmethod
    def export_to_markdown(tasks: List[Task]) -> str:
        """
        Export tasks to clean, professional Markdown format.
        Args:
            tasks: List of Task objects to export
        Returns:
            Well-formatted Markdown string
        """
        output = []

        output.append("# Programming Tasks Export\n")
        output.append(f"**Exported:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        output.append(f"**Total Tasks:** {len(tasks)}\n")
        output.append(f"Total tasks: {len(tasks)}\n")
        output.append("---\n")

        for i, task in enumerate(tasks, 1):
            output.extend(ExportService._format_task(task, i))

        return "\n".join(output)

    @staticmethod
    def _create_code_box(content: str, border_color: str) -> "Table":
        """Create a styled code box table for PDF export."""
        from reportlab.lib.units import inch
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib.colors import HexColor

        box = Table(
            [[content]],
            colWidths=[5.8 * inch]
        )

        box.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f8fafc')),
            ('BOX', (0, 0), (-1, -1), 1, HexColor(border_color)),
            ('FONTNAME', (0, 0), (-1, -1), 'Courier'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('PADDING', (0, 0), (-1, -1), 10),
        ]))

        return box

    @staticmethod
    def _add_task_to_pdf_story(
        story: list,
        task: Task,
        index: int,
        styles: dict,
        is_last_task: bool
    ) -> None:
        """Add a single task's content to the PDF story."""
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.colors import HexColor
        from reportlab.lib import colors

        # Task title
        story.append(Paragraph(f"{index}. {task.title}", styles["TaskTitle"]))

        # Metadata table
        meta_data = [
            ["Language", task.language],
            ["Concept", task.concept],
            ["Difficulty", task.difficulty]
        ]

        if hasattr(task, "template_name") and task.template_name:
            meta_data.append(["Template", task.template_name])

        meta_table = Table(meta_data, colWidths=[1.4 * inch, 4.4 * inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), HexColor('#f8fafc')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))

        story.append(meta_table)
        story.append(Spacer(1, 10))

        # Description
        story.append(Paragraph("Description", styles["SectionTitle"]))
        story.append(Paragraph(task.description, styles["BodyText"]))

        # Optional sections
        if task.examples and task.examples.strip():
            story.append(Paragraph("Examples", styles["SectionTitle"]))
            examples_box = ExportService._create_code_box(task.examples, '#d1d5db')
            # Override padding for examples
            examples_box.setStyle(TableStyle([('PADDING', (0, 0), (-1, -1), 8)]))
            story.append(examples_box)

        if task.solution:
            story.append(Paragraph("Reference Solution", styles["SectionTitle"]))
            story.append(ExportService._create_code_box(task.solution, '#94a3b8'))

        if task.tests and task.tests.strip():
            story.append(Paragraph("Tests", styles["SectionTitle"]))
            story.append(ExportService._create_code_box(task.tests, '#94a3b8'))

        if hasattr(task, "is_validated") and task.is_validated:
            validation = task.validation_result or {}
            status = validation.get("status", "unknown")
            status_text = "PASSED" if status == "passed" else "FAILED"

            story.append(Spacer(1, 10))
            story.append(Paragraph(f"<b>Validation:</b> {status_text}", styles["BodyText"]))

        if not is_last_task:
            story.append(PageBreak())

    @staticmethod
    def _create_pdf_styles(base_styles) -> dict:
        """Create and return custom PDF styles."""
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.colors import HexColor

        # Create a dictionary with the base styles we need
        styles = {
            "Normal": base_styles["Normal"],
            "Title": base_styles["Title"],
            "Heading1": base_styles["Heading1"],
            "Heading2": base_styles["Heading2"],
            "BodyText": base_styles["BodyText"],
        }

        styles["MainTitle"] = ParagraphStyle(
            name="MainTitle",
            parent=base_styles["Title"],
            fontSize=28,
            leading=34,
            alignment=TA_CENTER,
            textColor=HexColor("#1a365d"),
            spaceAfter=20
        )

        styles["Subtitle"] = ParagraphStyle(
            name="Subtitle",
            parent=base_styles["Normal"],
            fontSize=14,
            alignment=TA_CENTER,
            textColor=HexColor("#64748b"),
            spaceAfter=12
        )

        styles["TaskTitle"] = ParagraphStyle(
            name="TaskTitle",
            parent=base_styles["Heading1"],
            fontSize=18,
            leading=22,
            textColor=HexColor("#1e40af"),
            spaceAfter=12
        )

        styles["SectionTitle"] = ParagraphStyle(
            name="SectionTitle",
            parent=base_styles["Heading2"],
            fontSize=13,
            textColor=HexColor("#334155"),
            spaceBefore=10,
            spaceAfter=6
        )

        return styles

    @staticmethod
    def _add_page_number(canvas, doc):
        """Add page numbers to PDF."""
        from reportlab.lib.units import inch

        canvas.saveState()
        canvas.setFont("Helvetica", 9)

        page_num = canvas.getPageNumber()

        if page_num > 1:
            canvas.drawRightString(7.3 * inch, 0.5 * inch, f"{page_num}")

        canvas.restoreState()

    @staticmethod
    def export_to_pdf(tasks: List[Task]) -> bytes:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.colors import HexColor
        import io
        from datetime import datetime

        buffer = io.BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.8 * inch,
            leftMargin=0.8 * inch,
            topMargin=0.8 * inch,
            bottomMargin=0.8 * inch
        )

        base_styles = getSampleStyleSheet()
        styles = ExportService._create_pdf_styles(base_styles)

        story = []

        # Title page content
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph("Programming Tasks Export", styles["MainTitle"]))
        story.append(Paragraph("EduCode Platform", styles["Subtitle"]))
        story.append(Spacer(1, 0.4 * inch))
        story.append(Paragraph(
            datetime.now().strftime("Generated on %d %B %Y at %H:%M"),
            styles["Subtitle"]
        ))
        story.append(Paragraph(f"Total Tasks: <b>{len(tasks)}</b>", styles["Subtitle"]))
        story.append(Spacer(1, 0.8 * inch))

        # Decorative line
        line = Table([[""]], colWidths=[5.8 * inch], rowHeights=[2])
        line.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), HexColor('#1a365d'))]))
        story.append(line)
        story.append(PageBreak())

        # Add tasks
        for i, task in enumerate(tasks, 1):
            is_last_task = (i == len(tasks))
            ExportService._add_task_to_pdf_story(story, task, i, styles, is_last_task)

        # Build PDF
        doc.build(story, onFirstPage=ExportService._add_page_number, onLaterPages=ExportService._add_page_number)

        pdf_bytes = buffer.getvalue()
        buffer.close()

        return pdf_bytes

    
    @staticmethod
    def export_to_moodle_xml(tasks: List[Task]) -> str:
        import xml.etree.ElementTree as ET
        from datetime import datetime
        from html import escape

        root = ET.Element("quiz")

        for task in tasks:
            question = ET.SubElement(root, "question", {"type": "essay"})

            name = ET.SubElement(question, "name")
            ET.SubElement(name, "text").text = task.title

            questiontext = ET.SubElement(
                question,
                "questiontext",
                {"format": "html"}
            )

            qtext = ET.SubElement(questiontext, "text")

            html_content = f"""
    <div>
    <h3>{escape(task.title)}</h3>

    <p>
    <strong>Language:</strong> {escape(task.language)}<br/>
    <strong>Concept:</strong> {escape(task.concept)}<br/>
    <strong>Difficulty:</strong> {escape(task.difficulty)}
    </p>

    <h4>Task Description</h4>
    <p>{escape(task.description)}</p>
    """

            if task.examples and task.examples.strip():
                html_content += f"""
    <h4>Examples</h4>
    <pre>{escape(task.examples)}</pre>
    """

            if task.solution:
                html_content += f"""
    <h4>Reference Solution</h4>
    <pre>{escape(task.solution)}</pre>
    """

            if task.tests and task.tests.strip():
                html_content += f"""
    <h4>Tests</h4>
    <pre>{escape(task.tests)}</pre>
    """

            html_content += "</div>"

            qtext.text = html_content

            ET.SubElement(question, "defaultgrade").text = "10"
            ET.SubElement(question, "penalty").text = "0"
            ET.SubElement(question, "hidden").text = "0"

            if task.solution:
                generalfeedback = ET.SubElement(
                    question,
                    "generalfeedback"
                )

                fb_text = ET.SubElement(generalfeedback, "text")
                fb_text.text = f"""
    <h4>Reference Solution:</h4>
    <pre>{escape(task.solution)}</pre>
    """

        xml_str = ET.tostring(
            root,
            encoding="unicode",
            method="xml"
        )

        header = f"""<?xml version="1.0" encoding="UTF-8"?>
    <!-- Moodle XML Export from EduCode -->
    <!-- Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} -->
    <!-- Total tasks: {len(tasks)} -->
    """

        return header + xml_str
    
    @staticmethod
    def get_export_filename(format: str, prefix: str = "edocode_tasks") -> str:
        """
        Generate professional filename for different export formats.
        """
        extensions = {
            "pdf": "pdf",
            "markdown": "md",
            "moodle_xml": "xml",
            "docx": "docx",
            "json": "json"
        }

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        ext = extensions.get(format.lower(), "txt")
        
        return f"{prefix}_{timestamp}.{ext}"
    
    @staticmethod
    def _setup_document_margins(document) -> None:
        """Setup document margins and header/footer settings."""
        from docx.shared import Inches

        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.different_first_page_header_footer = True

    @staticmethod
    def _add_title_page(document, tasks) -> None:
        """Add elegant title page to the document."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime

        for _ in range(5):
            document.add_paragraph()

        title = document.add_heading("Programming Tasks Export", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(42)
        title_run.font.bold = True

        subtitle = document.add_paragraph("EduCode Platform", style='Subtitle')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(18)
        subtitle.runs[0].font.italic = True

        document.add_paragraph()

        info = document.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f"Exported: {datetime.now().strftime('%d %B %Y at %H:%M:%S')}\n")
        run.font.size = Pt(13)
        run.font.italic = True

        run = info.add_run(f"Total tasks: {len(tasks)}")
        run.font.size = Pt(13)

        document.add_paragraph("═" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _add_task_metadata(document, task) -> None:
        """Add task metadata to the document."""
        meta_p = document.add_paragraph()
        meta_p.add_run("Language: ").bold = True
        meta_p.add_run(f"{task.language}  ")
        meta_p.add_run("Concept: ").bold = True
        meta_p.add_run(f"{task.concept}  ")
        meta_p.add_run("Difficulty: ").bold = True
        meta_p.add_run(task.difficulty)

        if hasattr(task, 'template_name') and task.template_name:
            meta_p = document.add_paragraph()
            meta_p.add_run("Template: ").bold = True
            meta_p.add_run(task.template_name)

    @staticmethod
    def _add_task_section(document, heading: str, content: str, 
                          is_code: bool = False, style: str = None) -> None:
        """Add a section with optional code formatting to the document."""
        from docx.shared import Pt

        document.add_heading(heading, level=2)
        
        if is_code:
            p = document.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            p = document.add_paragraph(content)
            if style and len(content) < 600:
                p.style = style

    @staticmethod
    def _add_task_validation(document, task) -> None:
        """Add validation status to the document."""
        validation = task.validation_result or {}
        status = validation.get('status', 'unknown')
        status_text = "✅ PASSED" if status == 'passed' else "❌ FAILED"
        p = document.add_paragraph()
        p.add_run("Validation: ").bold = True
        p.add_run(status_text)

    @staticmethod
    def _add_task_to_docx(document, task, index: int, is_last: bool) -> None:
        """Add a single task's content to the DOCX document."""
        from docx.shared import Pt

        document.add_heading(f"{index}. {task.title}", level=1)

        ExportService._add_task_metadata(document, task)

        document.add_paragraph()

        document.add_heading("Description", level=2)
        document.add_paragraph(task.description)

        if task.examples and task.examples.strip():
            ExportService._add_task_section(
                document, "Examples", task.examples, 
                is_code=False, style='Intense Quote'
            )

        if task.solution:
            ExportService._add_task_section(
                document, "Reference Solution", task.solution, 
                is_code=True
            )

        if task.tests and task.tests.strip():
            ExportService._add_task_section(
                document, "Tests", task.tests, 
                is_code=True
            )

        if hasattr(task, 'is_validated') and task.is_validated:
            ExportService._add_task_validation(document, task)

        if not is_last:
            document.add_page_break()

    @staticmethod
    def _setup_docx_footers(document) -> None:
        """Setup page numbering in document footers."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        for section in document.sections:
            footer = section.footer
            for para in footer.paragraphs:
                para.clear()
            
            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run("").font.size = Pt(10)
            
            run = p.add_run()
            fld = run._element.makeelement(
                qn('w:fldSimple'), 
                attrib={qn('w:instr'): 'PAGE'}
            )
            run._element.append(fld)

    @staticmethod
    def export_to_docx(tasks: List[Task]) -> bytes:
        """
        Professional DOCX export with elegant title page and page numbering.
        """
        import io
        from docx import Document

        document = Document()

        ExportService._setup_document_margins(document)
        ExportService._add_title_page(document, tasks)
        document.add_page_break()

        for i, task in enumerate(tasks, 1):
            is_last = (i == len(tasks))
            ExportService._add_task_to_docx(document, task, i, is_last)

        ExportService._setup_docx_footers(document)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
